

from user_research_helper.campaign.config import config
from typing import List, Tuple
import shutil
import os
from user_research_helper.result_analysis.transcript_report_parsing import parse_transcript_report, create_segment_dataset_from_interview_dataset
from user_research_helper.result_analysis.data import SegmentAnswer
import json
from user_research_helper.result_analysis.segment_report_builder import create_excel_report
from user_research_helper.result_analysis.segment_report_parsing import parse_segment_report
from user_research_helper.result_analysis.result_analysis import analyze_question_across_segments
from user_research_helper.result_analysis.result_report_builder import create_result_report
from user_research_helper.result_analysis.quote_addition import add_quotes_from_excel
from user_research_helper.result_analysis.result_analysis import ResultAnalysis
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH




def process_analysis(
    root_dir: str = "data"
) -> None:
    """
    Main function to process interviews
    
    Args:
        root_dir: Root directory containing all project files
    """
    try:
        # Initialize configuration
        config.initialize(root_dir)
        if config.should_debug('verbose'):
            print(f"Initialized configuration:")
            print(f"  Root directory: {config.root_dir}")
            print(f"  Language: {config.language}")
            
        analysis_dir = config.get_path('analysis_dir')
        
        segment_dir=config.get_path('segment_analysis_dir')
        os.makedirs(segment_dir, exist_ok=True)
        
        # files used between steps
        segment_report_file = os.path.join(segment_dir, "segment_analysis_report.xlsx")
        question_synthesis_json_file = os.path.join(analysis_dir, "results.json")
        
        #########
        ## step 1 - segment summaries
        ########
        if (config.get_config('do_segment_summaries', False)):
            transcript_report_file = os.path.join(analysis_dir, "transcript_analysis_report.xlsx")  
            
            interview_dataset = parse_transcript_report(transcript_report_file)
            
            # dump interview dataset to json
            interview_dataset_json_file = os.path.join(segment_dir, "interview_dataset.json")
            with open(interview_dataset_json_file, 'w', encoding='utf-8') as f:
                json.dump(interview_dataset.model_dump(), f, ensure_ascii=False, indent=2)
                
            segment_dataset = create_segment_dataset_from_interview_dataset(interview_dataset)
            # dump segment dataset to json
            segment_dataset_json_file = os.path.join(segment_dir, "segment_dataset.json")
            with open(segment_dataset_json_file, 'w', encoding='utf-8') as f:
                json.dump(segment_dataset.model_dump(), f, ensure_ascii=False, indent=2)
            
            # For each segment, analyze answers if not already done
            from user_research_helper.result_analysis.answers_analysis import analyze_segment_answers
            for segment_name, segment_answers in segment_dataset.segments.items():
                segment_file = os.path.join(segment_dir, f"{segment_name}.json")
                if not os.path.exists(segment_file):
                    # Get question texts
                    question_texts = {q.id: q.text for q in segment_dataset.questions}
                    
                    # Analyze each answer in the segment
                    for question_id, answer in segment_answers.items():
                        analyze_segment_answers(answer, question_texts[question_id])
                        # Save progress after each answer
                        with open(segment_file, 'w', encoding='utf-8') as f:
                            json.dump({
                                "segment_name": segment_name,
                                "answers": {
                                    qid: answer.model_dump() 
                                    for qid, answer in segment_answers.items()
                                }
                            }, f, ensure_ascii=False, indent=2)
                        if config.should_debug('verbose'):
                            print(f"Segment {segment_name} analyzed and saved to {segment_file}")
                    else :
                        if config.should_debug('verbose'):
                            print(f"Segment {segment_name} already analyzed")
                        #load segment summary from json
                        with open(segment_file, 'r', encoding='utf-8') as f:
                            segment_summary = json.load(f)
                            segment_dataset.segments[segment_name] = {
                            qid: SegmentAnswer.model_validate({
                                **answer_data,
                                # Handle case where answer_summary is an object
                                "answer_summary": "\n".join(answer_data["answer_summary"].values()) 
                                if isinstance(answer_data["answer_summary"], dict) 
                                else answer_data["answer_summary"]
                            }) 
                            for qid, answer_data in segment_summary['answers'].items()
                        }
            
            # dump segment dataset to excel
            create_excel_report(segment_dataset, segment_report_file)
            if config.should_debug('verbose'):
                print(f"Segment report saved to {segment_report_file} ")
        
        ######
        ## step 2 - result analysis
        #######
        
        if (config.get_config('do_result_analysis', False)):
            if config.should_debug('verbose'):
                print(f"Make result analysis ")
            segment_dataset = parse_segment_report(segment_report_file)
            
            #dump segment dataset to json
            if (False):
                segment_dataset_json_file = os.path.join(analysis_dir, "segment_dataset_check.json")
                with open(segment_dataset_json_file, 'w', encoding='utf-8') as f:
                    json.dump(segment_dataset.model_dump(), f, ensure_ascii=False, indent=2)
                
            
            # for each question in the segment dataset, generate a synthesis of all segment summaries

        
            result_analysis_list = []
            for question in segment_dataset.questions:
                result_analysis = analyze_question_across_segments(segment_dataset, question.text, question.id)
                result_analysis_list.append(result_analysis)
                # Convert list of ResultAnalysis to list of dicts before JSON dump
                with open(question_synthesis_json_file, 'w', encoding='utf-8') as f:
                    json.dump([ra.model_dump() for ra in result_analysis_list], f, ensure_ascii=False, indent=2)       

            if config.should_debug('verbose'):
                print(f"Result analysis was saved to {question_synthesis_json_file}")
                
            result_report_file = os.path.join(analysis_dir, "result_report.xlsx")
            create_result_report(result_analysis_list, result_report_file)
         
            if config.should_debug('verbose'):
                print(f"Result report was saved to {result_report_file}")
         
         
        ######
        ## step 3 - add quotes
        #######
        
        if (config.get_config('do_add_quotes', False)):
            if config.should_debug('verbose'):
                print(f"Add quotes to results")
            transcript_report_file_quotes = os.path.join(analysis_dir, "transcript_analysis_report_quotes.xlsx") 
            question_synthesis_json_file_quotes = os.path.join(analysis_dir, "results_with_quotes.json")
            result_analysis_list = []
            with open(question_synthesis_json_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                result_analysis_list = [ResultAnalysis.model_validate({**ra, "quotes": ""}) for ra in data]
            

            
            result_analysis_list = add_quotes_from_excel(result_analysis_list, transcript_report_file_quotes)
            
            with open(question_synthesis_json_file_quotes, 'w', encoding='utf-8') as f:
                json.dump([ra.model_dump() for ra in result_analysis_list], f, ensure_ascii=False, indent=2)
            
            # Create Word document
            docx_file = os.path.join(analysis_dir, "results_with_quotes.docx")
            doc = Document()
            
            # Add title
            title = doc.add_heading('Analysis Results', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for ra in result_analysis_list:
                # Add question as heading level 2
                doc.add_heading(f"Q{ra.question_id} - {ra.question_text}", level=2)
                doc.add_paragraph()
                # Add analysis as normal paragraph
                doc.add_paragraph(ra.analysis)
                
                # Add quotes in italics if they exist
                if ra.quotes:
                    quote_para = doc.add_paragraph()
                    quote_run = quote_para.add_run(ra.quotes)
                    quote_run.italic = True
                
                # Add some space between sections
                doc.add_paragraph()
            
            # Save the document
            doc.save(docx_file)
                
            if config.should_debug('verbose'):
                print(f"Updated result_analysis_list with quotes was saved to {question_synthesis_json_file_quotes}")
                print(f"Word document was saved to {docx_file}")
            
    except (ValueError, FileNotFoundError, RuntimeError) as e:
        print(f"Error: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return

if __name__ == "__main__":
    import argparse
    import os
    from typing import List, Tuple
    
    parser = argparse.ArgumentParser(description='Process analysis of segmented transcript report')
    parser.add_argument('root_dir', default="demo", nargs='?', help='Root directory containing all project files')
    
    
    args = parser.parse_args()
    
    process_analysis(
        root_dir=args.root_dir
    )
